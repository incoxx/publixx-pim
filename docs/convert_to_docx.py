#!/usr/bin/env python3
"""
Convert verkaufsmappe.md to verkaufsmappe.docx with Dark Mode styling.
"""

import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Dark Mode Colors
BG_DARK = RGBColor(0x1E, 0x1E, 0x2E)       # Dark background
BG_SECTION = RGBColor(0x28, 0x28, 0x3C)     # Section background
TEXT_PRIMARY = RGBColor(0xCD, 0xD6, 0xF4)    # Light text
TEXT_SECONDARY = RGBColor(0xA6, 0xAD, 0xC8)  # Muted text
ACCENT_BLUE = RGBColor(0x89, 0xB4, 0xFA)     # Headings
ACCENT_LAVENDER = RGBColor(0xB4, 0xBE, 0xFE) # Sub-headings
ACCENT_GREEN = RGBColor(0xA6, 0xE3, 0xA1)    # Code/highlights
TABLE_HEADER_BG = "313244"                     # Table header bg
TABLE_ROW_BG = "1E1E2E"                       # Table row bg
TABLE_ALT_BG = "252536"                       # Table alt row bg
TABLE_BORDER = "45475A"                        # Table borders
CODE_BG = "181825"                             # Code block bg


def set_cell_bg(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_page_bg(doc, color_hex):
    """Set the page background color."""
    section = doc.sections[0]
    sectPr = section._sectPr
    bg = parse_xml(
        f'<w:background {nsdecls("w")} w:color="{color_hex}" w:themeColor="dark1"/>'
    )
    # Insert background as first child
    sectPr.insert(0, bg)
    # Also set displayBackgroundShape
    settings = doc.settings.element
    display_bg = parse_xml(
        f'<w:displayBackgroundShape {nsdecls("w")}/>'
    )
    settings.append(display_bg)


def create_style(doc, name, font_name, font_size, color, bold=False, space_before=0, space_after=6):
    """Create or get a paragraph style."""
    styles = doc.styles
    try:
        style = styles[name]
    except KeyError:
        style = styles.add_style(name, 1)  # WD_STYLE_TYPE.PARAGRAPH
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.color.rgb = color
    font.bold = bold
    pf = style.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    return style


def add_styled_paragraph(doc, text, style_name, alignment=None):
    """Add a paragraph with a named style."""
    p = doc.add_paragraph(text, style=style_name)
    if alignment is not None:
        p.alignment = alignment
    return p


def add_code_block(doc, code_text):
    """Add a code block with dark background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Add shading to paragraph
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CODE_BG}" w:val="clear"/>')
    pPr.append(shading)
    # Set left indent
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    for line in code_text.split('\n'):
        run = p.add_run(line + '\n')
        run.font.name = 'Courier New'
        run.font.size = Pt(8.5)
        run.font.color.rgb = ACCENT_GREEN
    return p


def add_table(doc, headers, rows):
    """Add a styled dark-mode table."""
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True

    # Style header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.font.color.rgb = TEXT_PRIMARY
        run.font.bold = True
        set_cell_bg(cell, TABLE_HEADER_BG)

    # Style data rows
    for row_idx, row_data in enumerate(rows):
        bg = TABLE_ALT_BG if row_idx % 2 == 0 else TABLE_ROW_BG
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= num_cols:
                break
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            # Handle bold markers
            clean_text = cell_text.strip()
            is_bold = clean_text.startswith('**') and clean_text.endswith('**')
            if is_bold:
                clean_text = clean_text[2:-2]
            run = p.add_run(clean_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(8.5)
            run.font.color.rgb = TEXT_PRIMARY
            if is_bold:
                run.font.bold = True
            set_cell_bg(cell, bg)

    # Set table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    return table


def parse_table(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (headers, rows, end_idx)."""
    headers = [h.strip() for h in lines[start_idx].strip('|').split('|')]
    # Skip separator line
    row_start = start_idx + 2
    rows = []
    idx = row_start
    while idx < len(lines) and lines[idx].strip().startswith('|'):
        row = [c.strip() for c in lines[idx].strip('|').split('|')]
        rows.append(row)
        idx += 1
    return headers, rows, idx


def process_inline(paragraph, text, default_color=TEXT_PRIMARY, default_size=10):
    """Process inline markdown (bold, code, links) and add runs to paragraph."""
    # Pattern for **bold**, `code`, and [text](url)
    pattern = r'(\*\*(.+?)\*\*|`([^`]+)`|\[([^\]]+)\]\(([^)]+)\)|([^*`\[]+))'
    for match in re.finditer(pattern, text):
        if match.group(2):  # Bold
            run = paragraph.add_run(match.group(2))
            run.font.name = 'Calibri'
            run.font.size = Pt(default_size)
            run.font.color.rgb = default_color
            run.font.bold = True
        elif match.group(3):  # Code
            run = paragraph.add_run(match.group(3))
            run.font.name = 'Courier New'
            run.font.size = Pt(default_size - 1)
            run.font.color.rgb = ACCENT_GREEN
        elif match.group(4):  # Link
            run = paragraph.add_run(f'{match.group(4)} ({match.group(5)})')
            run.font.name = 'Calibri'
            run.font.size = Pt(default_size)
            run.font.color.rgb = ACCENT_BLUE
            run.font.underline = True
        elif match.group(6):  # Plain text
            run = paragraph.add_run(match.group(6))
            run.font.name = 'Calibri'
            run.font.size = Pt(default_size)
            run.font.color.rgb = default_color


def convert_md_to_docx(md_path, docx_path):
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Set dark background
    set_page_bg(doc, "1E1E2E")

    # Create styles
    create_style(doc, 'DarkTitle', 'Calibri', 28, ACCENT_BLUE, bold=True, space_before=0, space_after=12)
    create_style(doc, 'DarkH2', 'Calibri', 18, ACCENT_BLUE, bold=True, space_before=24, space_after=8)
    create_style(doc, 'DarkH3', 'Calibri', 13, ACCENT_LAVENDER, bold=True, space_before=16, space_after=6)
    create_style(doc, 'DarkBody', 'Calibri', 10, TEXT_PRIMARY, space_before=2, space_after=4)
    create_style(doc, 'DarkBodyMuted', 'Calibri', 9, TEXT_SECONDARY, space_before=2, space_after=4)
    create_style(doc, 'DarkBullet', 'Calibri', 10, TEXT_PRIMARY, space_before=1, space_after=1)
    create_style(doc, 'DarkHR', 'Calibri', 2, RGBColor(0x45, 0x47, 0x5A), space_before=12, space_after=12)

    # Read markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_buffer = []
    skip_toc = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code block handling
        if stripped.startswith('```'):
            if in_code_block:
                # End code block
                add_code_block(doc, '\n'.join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
                code_buffer = []
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Skip TOC section
        if stripped == '## Inhaltsverzeichnis':
            skip_toc = True
            i += 1
            continue
        if skip_toc:
            if stripped.startswith('## ') and stripped != '## Inhaltsverzeichnis':
                skip_toc = False
            elif stripped == '---':
                skip_toc = False
                i += 1
                continue
            else:
                i += 1
                continue

        # Horizontal rule
        if stripped == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            # Add a thin line
            run = p.add_run('_' * 80)
            run.font.size = Pt(2)
            run.font.color.rgb = RGBColor(0x45, 0x47, 0x5A)
            i += 1
            continue

        # Title (# )
        if stripped.startswith('# ') and not stripped.startswith('## '):
            title_text = stripped[2:].strip()
            add_styled_paragraph(doc, title_text, 'DarkTitle', WD_ALIGN_PARAGRAPH.CENTER)
            i += 1
            continue

        # H2 (## )
        if stripped.startswith('## '):
            h2_text = stripped[3:].strip()
            add_styled_paragraph(doc, h2_text, 'DarkH2')
            i += 1
            continue

        # H3 (### )
        if stripped.startswith('### '):
            h3_text = stripped[4:].strip()
            add_styled_paragraph(doc, h3_text, 'DarkH3')
            i += 1
            continue

        # Table
        if stripped.startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
            headers, rows, end_idx = parse_table(lines, i)
            add_table(doc, headers, rows)
            i = end_idx
            continue

        # Bullet points
        if stripped.startswith('- '):
            bullet_text = stripped[2:]
            p = doc.add_paragraph(style='DarkBullet')
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            run = p.add_run('  ')
            run.font.color.rgb = ACCENT_BLUE
            run.font.size = Pt(10)
            process_inline(p, bullet_text)
            i += 1
            continue

        # Numbered list
        numbered_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if numbered_match:
            num = numbered_match.group(1)
            text = numbered_match.group(2)
            p = doc.add_paragraph(style='DarkBody')
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            run = p.add_run(f'{num}. ')
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.color.rgb = ACCENT_BLUE
            run.font.bold = True
            process_inline(p, text)
            i += 1
            continue

        # Bold paragraph (like **Schritt 1 — ...**)
        if stripped.startswith('**') and stripped.endswith('**'):
            p = doc.add_paragraph(style='DarkBody')
            process_inline(p, stripped)
            i += 1
            continue

        # Regular paragraph
        if stripped:
            p = doc.add_paragraph(style='DarkBody')
            process_inline(p, stripped)
            i += 1
            continue

        # Empty line
        i += 1

    # Save
    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == '__main__':
    convert_md_to_docx(
        '/home/user/publixx-pim/docs/verkaufsmappe.md',
        '/home/user/publixx-pim/docs/verkaufsmappe.docx'
    )
    convert_md_to_docx(
        '/home/user/publixx-pim/docs/guide-von-0-auf-100.md',
        '/home/user/publixx-pim/docs/guide-von-0-auf-100.docx'
    )
